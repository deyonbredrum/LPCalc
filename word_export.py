# -*- coding: utf-8 -*-
"""
Word 计算书导出模块 - 美化版
支持合并/分开导出，支持纯净导出
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_font(cell, text, font_name='宋体', font_size=10.5, bold=False,
                  color=None, alignment='left'):
    """设置单元格内容和字体"""
    cell.text = text
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_heading_font(paragraph, font_name='黑体', font_size=16, bold=True, color=(0, 0, 0)):
    """设置标题字体"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_background(cell, color_hex='D9E1F2'):
    """设置单元格背景色"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def build_basic_info_section(doc):
    """构建基本信息部分（共用）"""
    heading = doc.add_heading('一、基本信息', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    now = datetime.now()
    p = doc.add_paragraph()
    p.add_run('计算日期：').bold = True
    p.add_run(now.strftime("%Y年%m月%d日"))

    p = doc.add_paragraph()
    p.add_run('依据规范：').bold = True
    p.add_run('GB 50057-2010《建筑物防雷设计规范》')

    doc.add_paragraph()


def build_basic_info_section_with_ep(doc):
    """构建基本信息部分（含电子防护规范）"""
    heading = doc.add_heading('一、基本信息', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    now = datetime.now()
    p = doc.add_paragraph()
    p.add_run('计算日期：').bold = True
    p.add_run(now.strftime("%Y年%m月%d日"))

    p = doc.add_paragraph()
    p.add_run('依据规范：').bold = True
    p.add_run('GB 50057-2010《建筑物防雷设计规范》')

    p = doc.add_paragraph()
    p.add_run('          ').bold = True
    p.add_run('GB 50343-2012《建筑物电子信息系统防雷技术规范》')

    doc.add_paragraph()


def build_input_params_section(doc, building_params, c_factors=None, cable_params=None, include_ep=True):
    """
    构建输入参数部分

    参数:
        include_ep: 是否包含电子防护相关参数（C1~C6, 入户设施）
    """
    heading = doc.add_heading('二、输入参数', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    # 2.1 建筑物参数（始终包含）
    heading3 = doc.add_heading('2.1 建筑物参数', level=2)
    set_heading_font(heading3, '黑体', 14, True, (0, 0, 0))

    table1 = doc.add_table(rows=len(building_params) + 1, cols=2)
    table1.style = 'Light Grid Accent 1'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.columns[0].width = Cm(4.5)
    table1.columns[1].width = Cm(11)

    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '参数名称'
    hdr_cells[1].text = '数值'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.name = '黑体'
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_cell_background(cell, 'D9E1F2')

    for i, (key, value) in enumerate(building_params.items(), 1):
        row = table1.rows[i]
        set_cell_font(row.cells[0], key, '宋体', 10.5, False)
        set_cell_font(row.cells[1], str(value), '宋体', 10.5, False)

    doc.add_paragraph()

    # 如果不包含电子防护，直接返回
    if not include_ep:
        return

    # 2.2 电子信息系统防护因子
    if c_factors:
        heading3 = doc.add_heading('2.2 电子信息系统防护因子', level=2)
        set_heading_font(heading3, '黑体', 14, True, (0, 0, 0))

        table2 = doc.add_table(rows=len(c_factors) + 1, cols=3)
        table2.style = 'Light Grid Accent 1'
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.columns[0].width = Cm(4)
        table2.columns[1].width = Cm(8)
        table2.columns[2].width = Cm(3.5)

        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = '因子'
        hdr_cells[1].text = '选择项'
        hdr_cells[2].text = '取值'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.name = '黑体'
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            set_cell_background(cell, 'D9E1F2')

        factor_names = {
            'C1': '建筑物材料结构',
            'C2': '信息系统重要程度',
            'C3': '设备耐冲击能力',
            'C4': '雷电防护区',
            'C5': '雷击事故后果',
            'C6': '区域雷暴等级',
        }

        for i, (key, value) in enumerate(c_factors.items(), 1):
            row = table2.rows[i]
            set_cell_font(row.cells[0], f"{key}（{factor_names.get(key, '')}）", '宋体', 10.5, False)
            set_cell_font(row.cells[1], value.get('type', ''), '宋体', 10.5, False)
            set_cell_font(row.cells[2], str(value.get('val', '')), 'Arial', 11, True, (0, 51, 102))

        doc.add_paragraph()

    # 2.3 入户设施参数
    if cable_params:
        heading3 = doc.add_heading('2.3 入户设施参数', level=2)
        set_heading_font(heading3, '黑体', 14, True, (0, 0, 0))

        table3 = doc.add_table(rows=len(cable_params) + 1, cols=2)
        table3.style = 'Light Grid Accent 1'
        table3.alignment = WD_TABLE_ALIGNMENT.CENTER
        table3.columns[0].width = Cm(4.5)
        table3.columns[1].width = Cm(11)

        hdr_cells = table3.rows[0].cells
        hdr_cells[0].text = '参数名称'
        hdr_cells[1].text = '数值'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.name = '黑体'
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            set_cell_background(cell, 'D9E1F2')

        for i, (key, value) in enumerate(cable_params.items(), 1):
            row = table3.rows[i]
            set_cell_font(row.cells[0], key, '宋体', 10.5, False)
            set_cell_font(row.cells[1], str(value), '宋体', 10.5, False)

        doc.add_paragraph()


def build_lp_section(doc, lp_result, level_text, building_attr="", attr_type=""):
    """构建防雷等级计算部分"""
    heading = doc.add_heading('三、防雷等级计算', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    calc_steps = [
        ('年雷暴日 Td', f"{lp_result.get('Td', 0):.2f} d/a"),
        ('雷击大地密度 Ng', f"{lp_result.get('Ng', 0):.4f} 次/(km²·a)"),
        ('扩大宽度 D', f"{lp_result.get('D', 0):.3f} m"),
        ('等效面积 Ae', f"{lp_result.get('Ae', 0):.6f} km²"),
        ('校正系数 k', f"{lp_result.get('k', 0):.2f}"),
        ('年预计雷击次数 N', f"{lp_result.get('N', 0):.6f} 次/a"),
    ]

    table = doc.add_table(rows=len(calc_steps) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '计算项'
    hdr_cells[1].text = '结果'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.name = '黑体'
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_cell_background(cell, 'D9E1F2')

    for i, (name, value) in enumerate(calc_steps, 1):
        row = table.rows[i]
        set_cell_font(row.cells[0], name, '宋体', 10.5, False)
        set_cell_font(row.cells[1], str(value), 'Arial', 10.5, False)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('▶ 防雷等级判定结果：').bold = True
    run = p.add_run(level_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    run.font.size = Pt(12)
    doc.add_paragraph()


def build_ep_section(doc, ep_result, ep_level_text):
    """构建电子防护等级计算部分"""
    heading = doc.add_heading('三、电子信息系统雷电防护等级计算', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    ep_steps = [
        ('建筑物年预计雷击次数 N1', f"{ep_result.get('N1', 0):.6f} 次/a"),
        ("电源线缆截收面积 Ae1'", f"{ep_result.get('Ae1', 0):.6f} km²"),
        ("信号线缆截收面积 Ae2'", f"{ep_result.get('Ae2', 0):.6f} km²"),
        ('入户设施年预计雷击次数 N2', f"{ep_result.get('N2', 0):.6f} 次/a"),
        ('总年预计雷击次数 N', f"{ep_result.get('N', 0):.6f} 次/a"),
        ('C = C1+C2+C3+C4+C5+C6', f"{ep_result.get('C', 0):.2f}"),
        ('可接受最大年雷击次数 Nc', f"{ep_result.get('Nc', 0):.6f} 次/a"),
        ('拦截效率 E', f"{ep_result.get('E', 0):.4f}"),
    ]

    table = doc.add_table(rows=len(ep_steps) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '计算项'
    hdr_cells[1].text = '结果'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.name = '黑体'
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_cell_background(cell, 'D9E1F2')

    for i, (name, value) in enumerate(ep_steps, 1):
        row = table.rows[i]
        set_cell_font(row.cells[0], name, '宋体', 10.5, False)
        set_cell_font(row.cells[1], str(value), 'Arial', 10.5, False)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('▶ 电子信息系统防护等级判定结果：').bold = True
    run = p.add_run(ep_level_text)
    run.bold = True

    if 'A' in ep_level_text:
        run.font.color.rgb = RGBColor(192, 0, 0)
    elif 'B' in ep_level_text:
        run.font.color.rgb = RGBColor(255, 128, 0)
    elif 'C' in ep_level_text:
        run.font.color.rgb = RGBColor(0, 112, 192)
    else:
        run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(12)
    doc.add_paragraph()


def get_lp_conclusion_text(level_text, N, building_attr, attr_type):
    """
    根据防雷等级和建筑属性生成规范依据文本
    """
    # 规范条文映射
    if level_text == "一类":
        clause = "第3.0.2条"
        detail = "制造、使用或贮存爆炸危险物质，且年平均雷击次数大于0.05次/a"
    elif level_text == "二类":
        if attr_type in ["crowded", "important", "heritage", "office", "comm", "station", "airport"]:
            clause = "第3.0.3条"
            detail = "人员密集的公共建筑物或重要场所，且年平均雷击次数大于0.05次/a"
        else:
            clause = "第3.0.3条"
            detail = "一般民用建筑，年平均雷击次数大于0.25次/a"
    elif level_text == "三类":
        if attr_type in ["crowded", "important", "heritage", "office", "comm", "station", "airport"]:
            clause = "第3.0.4条"
            detail = "人员密集的公共建筑物或重要场所，年平均雷击次数在0.01~0.05次/a之间"
        else:
            clause = "第3.0.4条"
            detail = "一般民用建筑，年平均雷击次数在0.05~0.25次/a之间"
    else:
        clause = "第3.0.4条"
        detail = "年平均雷击次数小于0.01次/a，可不设防雷"

    return f"根据GB 50057-2010《建筑物防雷设计规范》{clause}规定，{detail}，应判定为{level_text}防雷建筑物。"


def get_ep_conclusion_text(ep_level_text, E):
    """
    根据电子防护等级生成规范依据文本
    """
    # 规范条文映射（GB 50343-2012 第5.4.1条）
    clause = "第5.4.1条"

    if ep_level_text == "A级":
        detail = "雷电防护等级为A级，拦截效率E > 0.98"
    elif ep_level_text == "B级":
        detail = "雷电防护等级为B级，拦截效率0.95 < E ≤ 0.98"
    elif ep_level_text == "C级":
        detail = "雷电防护等级为C级，拦截效率0.90 < E ≤ 0.95"
    elif ep_level_text == "D级":
        detail = "雷电防护等级为D级，拦截效率0.80 < E ≤ 0.90"
    else:
        detail = "雷电防护等级低于D级，拦截效率E ≤ 0.80，可不设防护"

    return f"根据GB 50343-2012《建筑物电子信息系统防雷技术规范》{clause}规定，{detail}，应判定为{ep_level_text}。"


def build_conclusion_section(doc, lp_result, level_text, ep_result=None, ep_level_text=None,
                             has_ep=False, building_attr="", attr_type=""):
    """
    构建结论部分（带规范依据）
    """
    heading = doc.add_heading('四、结论', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))

    N = lp_result.get('N', 0)

    # ---- 防雷等级结论（带规范依据） ----
    lp_conclusion = get_lp_conclusion_text(level_text, N, building_attr, attr_type)
    p = doc.add_paragraph()
    run = p.add_run('1. ')
    run.bold = True
    p.add_run(lp_conclusion)

    # 加粗关键结果
    # 找到并加粗"应判定为xxx"
    for run in p.runs:
        if '应判定为' in run.text:
            # 拆分并重新设置格式
            pass

    doc.add_paragraph()

    # ---- 电子防护等级结论（带规范依据） ----
    if has_ep and ep_result and ep_level_text:
        E = ep_result.get('E', 0)
        ep_conclusion = get_ep_conclusion_text(ep_level_text, E)
        p = doc.add_paragraph()
        run = p.add_run('2. ')
        run.bold = True
        p.add_run(ep_conclusion)

    doc.add_paragraph()

    # 页脚
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('注：本计算书由 LPCalc 工具自动生成，仅供参考。正式设计需经专业审核。')
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def export_combined_report(lp_result, ep_result, building_params,
                            cable_params, c_factors, level_text, ep_level_text,
                            has_ep=True, building_attr="", attr_type=""):
    """
    导出合并版计算书
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 主标题（根据是否包含电子防护调整）
    if has_ep:
        title_text = '建筑物防雷及电子信息防护等级计算书'
    else:
        title_text = '建筑物防雷等级计算书'

    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_font(title, '黑体', 22, True, (0, 51, 102))
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 基本信息（根据是否包含电子防护调整）
    if has_ep:
        build_basic_info_section_with_ep(doc)
    else:
        build_basic_info_section(doc)

    # 输入参数（根据是否包含电子防护调整）
    build_input_params_section(doc, building_params, c_factors, cable_params, include_ep=has_ep)

    # 防雷等级计算（传递建筑物属性）
    build_lp_section(doc, lp_result, level_text, building_attr, attr_type)

    # 电子防护等级计算
    if has_ep and ep_result:
        build_ep_section(doc, ep_result, ep_level_text)

    # 结论（带规范依据）
    build_conclusion_section(doc, lp_result, level_text, ep_result, ep_level_text,
                             has_ep, building_attr, attr_type)

    file_buffer = io.BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer


def export_separate_reports(lp_result, ep_result, building_params,
                            cable_params, c_factors, level_text, ep_level_text):
    """
    导出分开版计算书（防雷和电子防护各一个文件）
    返回: (lp_file_buffer, ep_file_buffer)
    """
    # ========== 防雷等级计算书（纯净版，不含电子防护） ==========
    doc_lp = Document()
    style = doc_lp.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc_lp.add_heading('建筑物防雷等级计算书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_font(title, '黑体', 22, True, (0, 51, 102))
    doc_lp.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_lp.add_paragraph()

    # 基本信息（只含GB 50057）
    build_basic_info_section(doc_lp)

    # 输入参数（不含电子防护相关）
    build_input_params_section(doc_lp, building_params, include_ep=False)

    # 防雷等级计算
    build_lp_section(doc_lp, lp_result, level_text)

    # 结论（只含防雷）
    heading = doc_lp.add_heading('四、结论', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))
    p = doc_lp.add_paragraph()
    p.add_run('根据 GB 50057-2010 计算，该建筑物的防雷等级为：').bold = True
    run = p.add_run(f' {level_text}')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)
    run.font.size = Pt(12)

    doc_lp.add_paragraph()
    doc_lp.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc_lp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('注：本计算书由 LPCalc 工具自动生成，仅供参考。正式设计需经专业审核。')
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lp_buffer = io.BytesIO()
    doc_lp.save(lp_buffer)
    lp_buffer.seek(0)

    # ========== 电子防护等级计算书 ==========
    doc_ep = Document()
    style = doc_ep.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc_ep.add_heading('电子信息系统雷电防护等级计算书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_font(title, '黑体', 22, True, (0, 51, 102))
    doc_ep.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_ep.add_paragraph()

    # 基本信息（含GB 50343）
    build_basic_info_section_with_ep(doc_ep)

    # 输入参数（含电子防护相关）
    build_input_params_section(doc_ep, building_params, c_factors, cable_params, include_ep=True)

    if ep_result:
        build_ep_section(doc_ep, ep_result, ep_level_text)
    else:
        doc_ep.add_paragraph('电子信息系统防护等级：未计算')

    # 结论（只含电子防护）
    heading = doc_ep.add_heading('四、结论', level=1)
    set_heading_font(heading, '黑体', 16, True, (0, 51, 102))
    if ep_result:
        p = doc_ep.add_paragraph()
        p.add_run('根据 GB 50343-2012 计算，该建筑物的电子信息系统雷电防护等级为：').bold = True
        run = p.add_run(f' {ep_level_text}')
        run.bold = True
        if 'A' in ep_level_text:
            run.font.color.rgb = RGBColor(192, 0, 0)
        elif 'B' in ep_level_text:
            run.font.color.rgb = RGBColor(255, 128, 0)
        elif 'C' in ep_level_text:
            run.font.color.rgb = RGBColor(0, 112, 192)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)
        run.font.size = Pt(12)
    else:
        doc_ep.add_paragraph('电子信息系统防护等级：未计算')

    doc_ep.add_paragraph()
    doc_ep.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc_ep.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('注：本计算书由 LPCalc 工具自动生成，仅供参考。正式设计需经专业审核。')
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    ep_buffer = io.BytesIO()
    doc_ep.save(ep_buffer)
    ep_buffer.seek(0)

    return lp_buffer, ep_buffer


def export_calculation_report(lp_result, ep_result, building_params,
                               cable_params, c_factors, level_text, ep_level_text,
                               export_mode='combined', has_ep=True,
                               building_attr="", attr_type=""):
    """
    导出 Word 计算书（统一入口）
    """
    if export_mode == 'separate':
        return export_separate_reports(lp_result, ep_result, building_params,
                                        cable_params, c_factors, level_text, ep_level_text,
                                        building_attr, attr_type)
    else:
        return export_combined_report(lp_result, ep_result, building_params,
                                       cable_params, c_factors, level_text, ep_level_text,
                                       has_ep, building_attr, attr_type)